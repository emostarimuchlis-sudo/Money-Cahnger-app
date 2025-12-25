"""
User Manual Generator for Mulia Bali Valuta Money Changer Application
Generates comprehensive user manual in DOCX and PDF formats
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import os
from datetime import datetime

def create_user_manual_docx(output_path: str):
    """Create comprehensive user manual in DOCX format"""
    doc = Document()
    
    # Title
    title = doc.add_heading('PETUNJUK TEKNIS PENGGUNAAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('APLIKASI MULIA BALI VALUTA', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Money Changer Management System', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version info
    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info.add_run(f'\nVersi 1.0 | {datetime.now().strftime("%B %Y")}\n\n')
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('DAFTAR ISI', level=1)
    toc_items = [
        ('1. Pendahuluan', '3'),
        ('2. Login dan Logout', '4'),
        ('3. Dashboard', '5'),
        ('4. Manajemen Transaksi', '6'),
        ('5. Data Nasabah', '10'),
        ('6. Buku Kas', '12'),
        ('7. Mutasi Valas', '14'),
        ('8. Laporan', '16'),
        ('9. Pengaturan', '18'),
        ('10. Backup dan Export Data', '20'),
        ('11. FAQ dan Troubleshooting', '22'),
    ]
    
    table = doc.add_table(rows=len(toc_items), cols=2)
    for i, (item, page) in enumerate(toc_items):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = page
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # Chapter 1: Introduction
    doc.add_heading('1. PENDAHULUAN', level=1)
    
    doc.add_heading('1.1 Tentang Aplikasi', level=2)
    doc.add_paragraph(
        'Mulia Bali Valuta (MBA) Money Changer adalah aplikasi manajemen transaksi '
        'penukaran valuta asing yang dirancang khusus untuk membantu operasional '
        'money changer. Aplikasi ini mencakup fitur lengkap mulai dari pencatatan '
        'transaksi, manajemen nasabah, pelaporan, hingga audit trail.'
    )
    
    doc.add_heading('1.2 Fitur Utama', level=2)
    features = [
        'Pencatatan transaksi jual/beli valuta asing',
        'Multi-currency transaction (transaksi lebih dari 1 mata uang sekaligus)',
        'Manajemen data nasabah (Perorangan & Badan Usaha)',
        'Buku Kas harian dengan saldo otomatis',
        'Mutasi Valas per mata uang',
        'Laporan SIPESAT untuk pelaporan regulasi',
        'Export data ke Excel dan PDF',
        'Cetak nota/struk transaksi',
        'Multi-cabang dengan role-based access control',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.3 Role Pengguna', level=2)
    roles = [
        ('Admin', 'Akses penuh ke semua fitur, termasuk pengaturan sistem dan manajemen user'),
        ('Kasir', 'Dapat melihat semua transaksi di cabang, mengelola nasabah, dan membuat laporan'),
        ('Teller', 'Hanya dapat melihat dan membuat transaksi sendiri'),
    ]
    for role, desc in roles:
        p = doc.add_paragraph()
        p.add_run(f'{role}: ').bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # Chapter 2: Login
    doc.add_heading('2. LOGIN DAN LOGOUT', level=1)
    
    doc.add_heading('2.1 Cara Login', level=2)
    login_steps = [
        'Buka aplikasi melalui browser (Chrome/Firefox/Edge)',
        'Masukkan Email yang terdaftar',
        'Masukkan Password',
        'Klik tombol "Masuk"',
        'Jika berhasil, Anda akan diarahkan ke halaman Dashboard',
    ]
    for i, step in enumerate(login_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('2.2 Cara Logout', level=2)
    doc.add_paragraph('Klik menu "Keluar" di sidebar kiri bawah untuk logout dari aplikasi.')
    
    doc.add_heading('2.3 Lupa Password', level=2)
    doc.add_paragraph('Hubungi Administrator untuk reset password akun Anda.')
    
    doc.add_page_break()
    
    # Chapter 3: Dashboard
    doc.add_heading('3. DASHBOARD', level=1)
    
    doc.add_paragraph(
        'Dashboard menampilkan ringkasan operasional harian yang meliputi:'
    )
    
    dashboard_items = [
        'Transaksi Hari Ini - Jumlah total transaksi yang dilakukan hari ini',
        'Nasabah Aktif - Total nasabah yang terdaftar dan aktif',
        'Total Pembelian - Total nilai pembelian valas hari ini (dalam IDR)',
        'Total Penjualan - Total nilai penjualan valas hari ini (dalam IDR)',
        'Grafik Transaksi - Visualisasi tren transaksi harian',
        'Transaksi Terbaru - Daftar 5 transaksi terakhir',
    ]
    for item in dashboard_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Chapter 4: Transactions
    doc.add_heading('4. MANAJEMEN TRANSAKSI', level=1)
    
    doc.add_heading('4.1 Halaman Transaksi', level=2)
    doc.add_paragraph(
        'Halaman transaksi menampilkan daftar transaksi per tanggal. '
        'Gunakan navigasi tanggal untuk berpindah antar hari.'
    )
    
    doc.add_heading('4.2 Membuat Transaksi Baru (Single Currency)', level=2)
    create_steps = [
        'Klik tombol "+ Transaksi Baru"',
        'Cari dan pilih nasabah dari dropdown (ketik untuk mencari)',
        'Jika nasabah baru, klik "+ Nasabah Baru" untuk mendaftarkan',
        'Pilih Tipe Transaksi: Jual atau Beli',
        'Pilih Mata Uang yang akan ditransaksikan',
        'Masukkan Jumlah valas',
        'Masukkan Kurs (rate) yang berlaku',
        'Total IDR akan dihitung otomatis',
        'Isi No. Voucher jika ada (opsional)',
        'Minta nasabah untuk tanda tangan di kolom Tanda Tangan',
        'Klik "Simpan" atau "Cetak & Simpan" untuk menyimpan transaksi',
    ]
    for i, step in enumerate(create_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('4.3 Membuat Transaksi Multi-Currency', level=2)
    doc.add_paragraph(
        'Untuk transaksi dengan lebih dari satu mata uang sekaligus:'
    )
    multi_steps = [
        'Centang opsi "Multi-Currency Transaction"',
        'Klik "+ Tambah Mata Uang" untuk menambah baris',
        'Isi detail setiap mata uang (tipe, mata uang, jumlah, kurs)',
        'Sistem akan menghitung total keseluruhan',
        'Simpan transaksi seperti biasa',
    ]
    for i, step in enumerate(multi_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('4.4 Cetak Nota/Struk', level=2)
    doc.add_paragraph(
        'Untuk mencetak nota transaksi:'
    )
    print_steps = [
        'Dari daftar transaksi, klik ikon mata (lihat detail)',
        'Klik tombol "Cetak" untuk mencetak nota',
        'Atau klik ikon printer dengan tanda tangan untuk mencetak dengan tanda tangan nasabah',
    ]
    for i, step in enumerate(print_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('4.5 Export Data Transaksi', level=2)
    doc.add_paragraph('Klik tombol "Excel" untuk export ke format spreadsheet.')
    doc.add_paragraph('Klik tombol "PDF" untuk export ke format PDF.')
    
    doc.add_page_break()
    
    # Chapter 5: Customers
    doc.add_heading('5. DATA NASABAH', level=1)
    
    doc.add_heading('5.1 Jenis Nasabah', level=2)
    customer_types = [
        ('Perorangan', 'Individu dengan identitas KTP/Paspor/SIM'),
        ('Badan Usaha', 'Perusahaan/organisasi dengan NPWP dan akta pendirian'),
    ]
    for ctype, desc in customer_types:
        p = doc.add_paragraph()
        p.add_run(f'{ctype}: ').bold = True
        p.add_run(desc)
    
    doc.add_heading('5.2 Menambah Nasabah Baru', level=2)
    add_customer_steps = [
        'Buka menu "Data Nasabah"',
        'Klik tombol "+ Nasabah Baru"',
        'Pilih Tipe Nasabah (Perorangan/Badan Usaha)',
        'Isi data identitas sesuai jenis nasabah',
        'Upload foto identitas jika diperlukan',
        'Klik "Simpan"',
    ]
    for i, step in enumerate(add_customer_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('5.3 Melihat Profil dan Riwayat Transaksi', level=2)
    doc.add_paragraph(
        'Klik ikon profil pada daftar nasabah untuk melihat detail lengkap '
        'termasuk riwayat transaksi dan total YTD (Year-to-Date).'
    )
    
    doc.add_page_break()
    
    # Chapter 6: Cash Book
    doc.add_heading('6. BUKU KAS', level=1)
    
    doc.add_heading('6.1 Fungsi Buku Kas', level=2)
    doc.add_paragraph(
        'Buku Kas mencatat semua aliran uang tunai (Rupiah) dari transaksi jual/beli valas.'
    )
    
    doc.add_heading('6.2 Jenis Entri', level=2)
    entry_types = [
        ('Debit', 'Uang masuk (dari penjualan valas ke nasabah)'),
        ('Kredit', 'Uang keluar (untuk pembelian valas dari nasabah)'),
    ]
    for etype, desc in entry_types:
        p = doc.add_paragraph()
        p.add_run(f'{etype}: ').bold = True
        p.add_run(desc)
    
    doc.add_heading('6.3 Navigasi Periode', level=2)
    doc.add_paragraph(
        'Gunakan tombol navigasi tanggal untuk melihat buku kas per hari. '
        'Saldo awal adalah saldo akhir dari hari sebelumnya.'
    )
    
    doc.add_heading('6.4 Menambah Entri Manual', level=2)
    doc.add_paragraph(
        'Untuk menambah entri kas non-transaksi (misalnya setor/tarik kas):'
    )
    manual_entry_steps = [
        'Klik tombol "+ Entri Baru"',
        'Pilih Tipe (Debit/Kredit)',
        'Masukkan Jumlah',
        'Isi Keterangan',
        'Klik "Simpan"',
    ]
    for i, step in enumerate(manual_entry_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # Chapter 7: Mutasi Valas
    doc.add_heading('7. MUTASI VALAS', level=1)
    
    doc.add_heading('7.1 Fungsi Mutasi Valas', level=2)
    doc.add_paragraph(
        'Mutasi Valas menampilkan pergerakan stok setiap mata uang asing '
        'termasuk pembelian, penjualan, dan saldo akhir.'
    )
    
    doc.add_heading('7.2 Informasi yang Ditampilkan', level=2)
    mutasi_info = [
        'Stock Awal - Saldo valas di awal periode',
        'Pembelian - Total valas yang dibeli dari nasabah',
        'Penjualan - Total valas yang dijual ke nasabah',
        'Stock Akhir - Saldo valas di akhir periode',
        'Rata-rata Kurs - Kurs rata-rata tertimbang',
        'Laba/Rugi - Selisih kurs yang menghasilkan keuntungan/kerugian',
    ]
    for info in mutasi_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_heading('7.3 Ringkasan Total Rupiah', level=2)
    doc.add_paragraph(
        'Di bagian bawah tabel terdapat ringkasan total dalam Rupiah untuk '
        'semua mata uang yang ditampilkan.'
    )
    
    doc.add_page_break()
    
    # Chapter 8: Reports
    doc.add_heading('8. LAPORAN', level=1)
    
    doc.add_heading('8.1 Laporan SIPESAT', level=2)
    doc.add_paragraph(
        'SIPESAT adalah laporan untuk kepatuhan regulasi Bank Indonesia. '
        'Laporan ini menampilkan data transaksi per nasabah dalam periode triwulan.'
    )
    
    doc.add_heading('8.2 Periode SIPESAT', level=2)
    sipesat_periods = [
        'Q1: Januari - Maret',
        'Q2: April - Juni',
        'Q3: Juli - September',
        'Q4: Oktober - Desember',
    ]
    for period in sipesat_periods:
        doc.add_paragraph(period, style='List Bullet')
    
    doc.add_heading('8.3 Mengunci Periode', level=2)
    doc.add_paragraph(
        'Setelah periode berakhir, Admin dapat mengunci laporan untuk '
        'mencegah perubahan data. Data yang sudah dikunci tidak dapat diubah.'
    )
    
    doc.add_page_break()
    
    # Chapter 9: Settings
    doc.add_heading('9. PENGATURAN (Khusus Admin)', level=1)
    
    doc.add_heading('9.1 Manajemen Pengguna', level=2)
    doc.add_paragraph(
        'Admin dapat menambah, mengedit, dan menonaktifkan pengguna sistem.'
    )
    
    doc.add_heading('9.2 Manajemen Cabang', level=2)
    doc.add_paragraph(
        'Mengelola daftar cabang termasuk saldo awal kas dan stok valas.'
    )
    
    doc.add_heading('9.3 Manajemen Mata Uang', level=2)
    doc.add_paragraph(
        'Menambah atau mengedit daftar mata uang yang dapat ditransaksikan.'
    )
    
    doc.add_heading('9.4 Pengaturan Perusahaan', level=2)
    doc.add_paragraph(
        'Mengatur informasi perusahaan yang akan tampil di nota/struk.'
    )
    
    doc.add_heading('9.5 Log Aktivitas', level=2)
    doc.add_paragraph(
        'Melihat riwayat aktivitas pengguna untuk keperluan audit.'
    )
    
    doc.add_page_break()
    
    # Chapter 10: Backup & Export
    doc.add_heading('10. BACKUP DAN EXPORT DATA', level=1)
    
    doc.add_heading('10.1 Export ke Excel', level=2)
    doc.add_paragraph(
        'Hampir semua halaman menyediakan tombol "Excel" untuk mengexport '
        'data ke format spreadsheet (.xlsx). Data ini dapat dibuka di '
        'Microsoft Excel atau Google Sheets.'
    )
    
    doc.add_heading('10.2 Export ke PDF', level=2)
    doc.add_paragraph(
        'Tombol "PDF" akan mengexport data dalam format PDF yang siap cetak.'
    )
    
    doc.add_heading('10.3 Cetak Laporan', level=2)
    doc.add_paragraph(
        'Tombol "Cetak" akan membuka dialog print browser untuk mencetak '
        'langsung ke printer.'
    )
    
    doc.add_heading('10.4 API Endpoint untuk Backup', level=2)
    doc.add_paragraph(
        'Untuk keperluan backup database, gunakan endpoint berikut (khusus Admin):'
    )
    api_endpoints = [
        'GET /api/transactions - Mengambil semua data transaksi',
        'GET /api/customers - Mengambil semua data nasabah',
        'GET /api/cashbook - Mengambil data buku kas',
        'GET /api/mutasi-valas/calculate - Mengambil data mutasi valas',
    ]
    for endpoint in api_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_page_break()
    
    # Chapter 11: FAQ
    doc.add_heading('11. FAQ DAN TROUBLESHOOTING', level=1)
    
    faqs = [
        ('Tidak bisa login', 
         'Pastikan email dan password benar. Jika lupa password, hubungi Admin.'),
        ('Transaksi tidak muncul di Buku Kas',
         'Pastikan transaksi sudah disimpan. Cek tanggal yang dipilih di Buku Kas.'),
        ('Gagal menyimpan transaksi',
         'Pastikan semua field wajib sudah diisi. Cek koneksi internet.'),
        ('Error "Service Temporarily Unavailable"',
         'Coba refresh halaman. Jika masih error, tunggu beberapa saat dan coba lagi.'),
        ('Tidak bisa cetak nota',
         'Pastikan browser mengizinkan pop-up. Cek pengaturan printer.'),
        ('Data tidak ter-update',
         'Tekan Ctrl+F5 untuk refresh paksa. Clear cache browser jika perlu.'),
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(f'Q: {q}').bold = True
        doc.add_paragraph(f'A: {a}')
        doc.add_paragraph()
    
    # Footer
    doc.add_page_break()
    doc.add_heading('KONTAK DUKUNGAN', level=1)
    doc.add_paragraph('Untuk bantuan teknis, hubungi:')
    doc.add_paragraph('Email: support@moztec.com')
    doc.add_paragraph('WhatsApp: +62 xxx-xxxx-xxxx')
    doc.add_paragraph()
    doc.add_paragraph(f'Dokumen ini dibuat pada: {datetime.now().strftime("%d %B %Y")}')
    
    # Save document
    doc.save(output_path)
    return output_path


def create_user_manual_pdf(output_path: str):
    """Create comprehensive user manual in PDF format"""
    doc = SimpleDocTemplate(output_path, pagesize=A4, 
                           rightMargin=72, leftMargin=72, 
                           topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=12
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=20
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        parent=styles['Heading2'],
        fontSize=13,
        spaceAfter=8,
        spaceBefore=14
    ))
    
    styles.add(ParagraphStyle(
        name='CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    ))
    
    styles.add(ParagraphStyle(
        name='CustomBullet',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        spaceAfter=4
    ))
    
    story = []
    
    # Title Page
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph('PETUNJUK TEKNIS PENGGUNAAN', styles['CustomTitle']))
    story.append(Paragraph('APLIKASI MULIA BALI VALUTA', styles['CustomTitle']))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph('Money Changer Management System', styles['CustomBody']))
    story.append(Spacer(1, inch))
    story.append(Paragraph(f'Versi 1.0 | {datetime.now().strftime("%B %Y")}', styles['CustomBody']))
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph('DAFTAR ISI', styles['CustomHeading1']))
    toc_data = [
        ['1. Pendahuluan', '3'],
        ['2. Login dan Logout', '4'],
        ['3. Dashboard', '5'],
        ['4. Manajemen Transaksi', '6'],
        ['5. Data Nasabah', '10'],
        ['6. Buku Kas', '12'],
        ['7. Mutasi Valas', '14'],
        ['8. Laporan', '16'],
        ['9. Pengaturan', '18'],
        ['10. Backup dan Export Data', '20'],
        ['11. FAQ dan Troubleshooting', '22'],
    ]
    toc_table = Table(toc_data, colWidths=[4.5*inch, 0.5*inch])
    toc_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(toc_table)
    story.append(PageBreak())
    
    # Chapter 1
    story.append(Paragraph('1. PENDAHULUAN', styles['CustomHeading1']))
    
    story.append(Paragraph('1.1 Tentang Aplikasi', styles['CustomHeading2']))
    story.append(Paragraph(
        'Mulia Bali Valuta (MBA) Money Changer adalah aplikasi manajemen transaksi '
        'penukaran valuta asing yang dirancang khusus untuk membantu operasional '
        'money changer. Aplikasi ini mencakup fitur lengkap mulai dari pencatatan '
        'transaksi, manajemen nasabah, pelaporan, hingga audit trail.',
        styles['CustomBody']
    ))
    
    story.append(Paragraph('1.2 Fitur Utama', styles['CustomHeading2']))
    features = [
        '• Pencatatan transaksi jual/beli valuta asing',
        '• Multi-currency transaction',
        '• Manajemen data nasabah (Perorangan & Badan Usaha)',
        '• Buku Kas harian dengan saldo otomatis',
        '• Mutasi Valas per mata uang',
        '• Laporan SIPESAT untuk pelaporan regulasi',
        '• Export data ke Excel dan PDF',
        '• Cetak nota/struk transaksi',
        '• Multi-cabang dengan role-based access control',
    ]
    for f in features:
        story.append(Paragraph(f, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    # Chapter 2
    story.append(Paragraph('2. LOGIN DAN LOGOUT', styles['CustomHeading1']))
    
    story.append(Paragraph('2.1 Cara Login', styles['CustomHeading2']))
    login_steps = [
        '1. Buka aplikasi melalui browser (Chrome/Firefox/Edge)',
        '2. Masukkan Email yang terdaftar',
        '3. Masukkan Password',
        '4. Klik tombol "Masuk"',
        '5. Jika berhasil, Anda akan diarahkan ke halaman Dashboard',
    ]
    for step in login_steps:
        story.append(Paragraph(step, styles['CustomBullet']))
    
    story.append(Paragraph('2.2 Cara Logout', styles['CustomHeading2']))
    story.append(Paragraph(
        'Klik menu "Keluar" di sidebar kiri bawah untuk logout dari aplikasi.',
        styles['CustomBody']
    ))
    
    story.append(PageBreak())
    
    # Chapter 3
    story.append(Paragraph('3. DASHBOARD', styles['CustomHeading1']))
    story.append(Paragraph(
        'Dashboard menampilkan ringkasan operasional harian:',
        styles['CustomBody']
    ))
    
    dashboard_items = [
        '• Transaksi Hari Ini - Jumlah total transaksi hari ini',
        '• Nasabah Aktif - Total nasabah yang terdaftar',
        '• Total Pembelian - Total nilai pembelian valas (IDR)',
        '• Total Penjualan - Total nilai penjualan valas (IDR)',
        '• Grafik Transaksi - Visualisasi tren transaksi',
        '• Transaksi Terbaru - Daftar 5 transaksi terakhir',
    ]
    for item in dashboard_items:
        story.append(Paragraph(item, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    # Chapter 4
    story.append(Paragraph('4. MANAJEMEN TRANSAKSI', styles['CustomHeading1']))
    
    story.append(Paragraph('4.1 Halaman Transaksi', styles['CustomHeading2']))
    story.append(Paragraph(
        'Halaman transaksi menampilkan daftar transaksi per tanggal. '
        'Gunakan navigasi tanggal untuk berpindah antar hari.',
        styles['CustomBody']
    ))
    
    story.append(Paragraph('4.2 Membuat Transaksi Baru', styles['CustomHeading2']))
    create_steps = [
        '1. Klik tombol "+ Transaksi Baru"',
        '2. Cari dan pilih nasabah dari dropdown',
        '3. Pilih Tipe Transaksi: Jual atau Beli',
        '4. Pilih Mata Uang',
        '5. Masukkan Jumlah valas',
        '6. Masukkan Kurs yang berlaku',
        '7. Minta nasabah tanda tangan',
        '8. Klik "Simpan" atau "Cetak & Simpan"',
    ]
    for step in create_steps:
        story.append(Paragraph(step, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    # Chapter 5
    story.append(Paragraph('5. DATA NASABAH', styles['CustomHeading1']))
    
    story.append(Paragraph('5.1 Jenis Nasabah', styles['CustomHeading2']))
    story.append(Paragraph(
        '<b>Perorangan:</b> Individu dengan identitas KTP/Paspor/SIM',
        styles['CustomBody']
    ))
    story.append(Paragraph(
        '<b>Badan Usaha:</b> Perusahaan dengan NPWP dan akta pendirian',
        styles['CustomBody']
    ))
    
    story.append(Paragraph('5.2 Menambah Nasabah Baru', styles['CustomHeading2']))
    add_steps = [
        '1. Buka menu "Data Nasabah"',
        '2. Klik tombol "+ Nasabah Baru"',
        '3. Pilih Tipe Nasabah',
        '4. Isi data identitas',
        '5. Klik "Simpan"',
    ]
    for step in add_steps:
        story.append(Paragraph(step, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    # Chapter 6
    story.append(Paragraph('6. BUKU KAS', styles['CustomHeading1']))
    
    story.append(Paragraph(
        'Buku Kas mencatat semua aliran uang tunai (Rupiah) dari transaksi.',
        styles['CustomBody']
    ))
    
    story.append(Paragraph('6.1 Jenis Entri', styles['CustomHeading2']))
    story.append(Paragraph(
        '<b>Debit:</b> Uang masuk (dari penjualan valas ke nasabah)',
        styles['CustomBody']
    ))
    story.append(Paragraph(
        '<b>Kredit:</b> Uang keluar (untuk pembelian valas dari nasabah)',
        styles['CustomBody']
    ))
    
    story.append(PageBreak())
    
    # Chapter 7
    story.append(Paragraph('7. MUTASI VALAS', styles['CustomHeading1']))
    
    story.append(Paragraph(
        'Mutasi Valas menampilkan pergerakan stok setiap mata uang asing.',
        styles['CustomBody']
    ))
    
    mutasi_info = [
        '• Stock Awal - Saldo valas di awal periode',
        '• Pembelian - Total valas yang dibeli dari nasabah',
        '• Penjualan - Total valas yang dijual ke nasabah',
        '• Stock Akhir - Saldo valas di akhir periode',
        '• Rata-rata Kurs - Kurs rata-rata tertimbang',
        '• Laba/Rugi - Selisih kurs',
    ]
    for info in mutasi_info:
        story.append(Paragraph(info, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    # Chapter 8-11 Summary
    story.append(Paragraph('8. LAPORAN', styles['CustomHeading1']))
    story.append(Paragraph(
        'Laporan SIPESAT untuk kepatuhan regulasi Bank Indonesia. '
        'Menampilkan data transaksi per nasabah dalam periode triwulan.',
        styles['CustomBody']
    ))
    
    story.append(PageBreak())
    
    story.append(Paragraph('9. PENGATURAN (Khusus Admin)', styles['CustomHeading1']))
    settings_items = [
        '• Manajemen Pengguna',
        '• Manajemen Cabang',
        '• Manajemen Mata Uang',
        '• Pengaturan Perusahaan',
        '• Log Aktivitas',
    ]
    for item in settings_items:
        story.append(Paragraph(item, styles['CustomBullet']))
    
    story.append(PageBreak())
    
    story.append(Paragraph('10. BACKUP DAN EXPORT DATA', styles['CustomHeading1']))
    story.append(Paragraph(
        'Gunakan tombol "Excel" atau "PDF" di setiap halaman untuk export data. '
        'Tombol "Cetak" untuk mencetak langsung.',
        styles['CustomBody']
    ))
    
    story.append(PageBreak())
    
    story.append(Paragraph('11. FAQ DAN TROUBLESHOOTING', styles['CustomHeading1']))
    
    faqs = [
        ('<b>Tidak bisa login?</b>', 'Pastikan email dan password benar.'),
        ('<b>Transaksi tidak muncul?</b>', 'Cek tanggal yang dipilih.'),
        ('<b>Error service unavailable?</b>', 'Refresh halaman atau tunggu beberapa saat.'),
    ]
    for q, a in faqs:
        story.append(Paragraph(q, styles['CustomBody']))
        story.append(Paragraph(a, styles['CustomBullet']))
        story.append(Spacer(1, 0.1*inch))
    
    story.append(PageBreak())
    
    # Contact
    story.append(Paragraph('KONTAK DUKUNGAN', styles['CustomHeading1']))
    story.append(Paragraph('Email: support@moztec.com', styles['CustomBody']))
    story.append(Paragraph('WhatsApp: +62 xxx-xxxx-xxxx', styles['CustomBody']))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(
        f'Dokumen ini dibuat pada: {datetime.now().strftime("%d %B %Y")}',
        styles['CustomBody']
    ))
    
    doc.build(story)
    return output_path


if __name__ == "__main__":
    # Generate both formats
    docx_path = "/app/backend/static/user_manual.docx"
    pdf_path = "/app/backend/static/user_manual.pdf"
    
    os.makedirs("/app/backend/static", exist_ok=True)
    
    create_user_manual_docx(docx_path)
    print(f"DOCX created: {docx_path}")
    
    create_user_manual_pdf(pdf_path)
    print(f"PDF created: {pdf_path}")
