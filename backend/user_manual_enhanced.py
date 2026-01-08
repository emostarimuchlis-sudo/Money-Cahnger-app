"""
Enhanced User Manual Generator with Detailed Step-by-Step Instructions
Including visual descriptions and annotations for each step
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from datetime import datetime
import os

def add_note_box(doc, text, title="Catatan"):
    """Add a styled note box to document"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    
    runner = p.add_run(f"📌 {title}: ")
    runner.bold = True
    runner.font.color.rgb = RGBColor(212, 175, 55)  # Gold color
    
    runner = p.add_run(text)
    runner.font.color.rgb = RGBColor(110, 231, 183)  # Green color

def add_warning_box(doc, text):
    """Add a styled warning box to document"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    
    runner = p.add_run("⚠️ Peringatan: ")
    runner.bold = True
    runner.font.color.rgb = RGBColor(251, 191, 36)  # Amber color
    
    runner = p.add_run(text)
    runner.font.color.rgb = RGBColor(254, 243, 199)

def add_step_with_screenshot(doc, step_number, title, description, visual_note=""):
    """Add a detailed step with visual description"""
    # Step title
    p = doc.add_paragraph()
    runner = p.add_run(f"Langkah {step_number}: {title}")
    runner.bold = True
    runner.font.size = Pt(12)
    runner.font.color.rgb = RGBColor(212, 175, 55)
    
    # Description
    doc.add_paragraph(description)
    
    # Visual note (what user will see on screen)
    if visual_note:
        p = doc.add_paragraph()
        runner = p.add_run("🖥️ Tampilan Layar: ")
        runner.bold = True
        runner.italic = True
        runner.font.color.rgb = RGBColor(110, 231, 183)
        
        runner = p.add_run(visual_note)
        runner.italic = True
        runner.font.color.rgb = RGBColor(209, 250, 229)
        
    doc.add_paragraph()  # Spacing

def create_enhanced_user_manual_docx(output_path: str):
    """Create enhanced user manual with detailed step-by-step instructions"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('PETUNJUK TEKNIS PENGGUNAAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('APLIKASI MULIA BALI VALUTA', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Money Changer Management System', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = version_info.add_run(f'\nVersi 2.0 - Edisi Bergambar\n{datetime.now().strftime("%B %Y")}\n\n')
    runner.italic = True
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('DAFTAR ISI', level=1)
    toc_items = [
        '1. Pendahuluan',
        '2. Login dan Logout',
        '3. Dashboard',
        '4. Manajemen Transaksi',
        '5. Data Nasabah',
        '6. Buku Kas',
        '7. Mutasi Valas',
        '8. Laporan',
        '9. Pengaturan',
        '10. Backup dan Export Data',
        '11. Tips dan Troubleshooting',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== CHAPTER 1: PENDAHULUAN ==========
    doc.add_heading('1. PENDAHULUAN', level=1)
    
    doc.add_heading('1.1 Tentang Aplikasi', level=2)
    doc.add_paragraph(
        'Mulia Bali Valuta (MBA) Money Changer adalah aplikasi manajemen transaksi '
        'penukaran valuta asing yang dirancang khusus untuk membantu operasional '
        'money changer. Aplikasi ini mencakup fitur lengkap mulai dari pencatatan '
        'transaksi, manajemen nasabah, pelaporan, hingga audit trail.'
    )
    
    doc.add_heading('1.2 Fitur Utama', level=2)
    features = [
        'Pencatatan transaksi jual/beli valuta asing',
        'Multi-currency transaction (transaksi lebih dari 1 mata uang sekaligus)',
        'Manajemen data nasabah (Perorangan & Badan Usaha)',
        'Buku Kas harian dengan saldo otomatis',
        'Mutasi Valas per mata uang',
        'Laporan SIPESAT untuk pelaporan regulasi',
        'Export data ke Excel dan PDF',
        'Cetak nota/struk transaksi dengan tanda tangan nasabah',
        'Multi-cabang dengan role-based access control',
        'Tools diagnostik untuk memastikan konsistensi data',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.3 Role Pengguna', level=2)
    roles = [
        ('Admin', 'Akses penuh ke semua fitur, termasuk pengaturan sistem, manajemen user, dan tools maintenance'),
        ('Kasir', 'Dapat melihat semua transaksi di cabang, mengelola nasabah, dan membuat laporan'),
        ('Teller', 'Hanya dapat melihat dan membuat transaksi sendiri'),
    ]
    for role, desc in roles:
        p = doc.add_paragraph()
        runner = p.add_run(f'{role}: ')
        runner.bold = True
        runner.font.color.rgb = RGBColor(212, 175, 55)
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========== CHAPTER 2: LOGIN ==========
    doc.add_heading('2. LOGIN DAN LOGOUT', level=1)
    
    doc.add_heading('2.1 Cara Login ke Aplikasi', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Aplikasi di Browser",
        "Buka browser (Chrome, Firefox, Edge, atau Safari) dan masukkan URL aplikasi yang diberikan oleh administrator.",
        "Anda akan melihat halaman login dengan logo MOZTEC di tengah layar. Ada form dengan 2 field: Email dan Password, serta tombol kuning 'Masuk' di bawahnya."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Masukkan Email",
        "Klik pada field 'Email' dan ketik alamat email yang telah terdaftar di sistem.",
        "Field Email akan aktif (border berwarna) saat diklik. Ketik email dengan format yang benar (contoh: user@company.com)"
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Masukkan Password",
        "Klik pada field 'Password' dan ketik password Anda. Password akan muncul sebagai titik-titik untuk keamanan.",
        "Ada ikon mata di sebelah kanan field Password. Klik untuk show/hide password jika perlu memeriksa."
    )
    
    add_step_with_screenshot(
        doc, 4,
        "Klik Tombol Masuk",
        "Setelah email dan password terisi, klik tombol kuning 'Masuk' untuk login.",
        "Tombol 'Masuk' berwarna kuning emas dan berada di bawah form. Jika berhasil, akan muncul notifikasi hijau 'Login berhasil!' di pojok kanan atas."
    )
    
    add_step_with_screenshot(
        doc, 5,
        "Masuk ke Dashboard",
        "Setelah login berhasil, Anda akan otomatis diarahkan ke halaman Dashboard.",
        "Dashboard menampilkan ringkasan hari ini: Transaksi, Pendapatan, Total Nasabah, Cabang Aktif, dan grafik transaksi terbaru."
    )
    
    add_note_box(doc, "Email dan password default untuk admin adalah admin@moztec.com / admin123. Segera ubah password setelah login pertama kali!")
    
    doc.add_heading('2.2 Cara Logout', level=2)
    add_step_with_screenshot(
        doc, 1,
        "Klik Menu Keluar",
        "Di sidebar kiri bawah, klik menu 'Keluar' untuk logout dari aplikasi.",
        "Menu 'Keluar' ada di paling bawah sidebar, di bawah nama pengguna. Icon panah keluar berwarna merah. Setelah klik, Anda akan kembali ke halaman login."
    )
    
    doc.add_page_break()
    
    # ========== CHAPTER 3: DASHBOARD ==========
    doc.add_heading('3. DASHBOARD', level=1)
    
    doc.add_paragraph(
        'Dashboard adalah halaman utama yang menampilkan ringkasan operasional harian. '
        'Halaman ini memberikan overview cepat tentang performa bisnis hari ini.'
    )
    
    doc.add_heading('3.1 Komponen Dashboard', level=2)
    
    dashboard_components = [
        ('Transaksi Hari Ini', 'Menampilkan jumlah total transaksi yang dilakukan hari ini. Icon: grafik naik berwarna hijau.'),
        ('Pendapatan Hari Ini', 'Total nilai transaksi hari ini dalam Rupiah. Icon: uang berwarna kuning.'),
        ('Total Nasabah', 'Jumlah nasabah yang terdaftar di sistem. Icon: orang berwarna biru.'),
        ('Cabang Aktif', 'Jumlah cabang yang aktif beroperasi. Icon: gedung berwarna ungu.'),
        ('Transaksi Terbaru', 'Tabel menampilkan 5 transaksi terakhir dengan detail nomor transaksi, nasabah, mata uang, tipe, dan total.'),
    ]
    
    for component, desc in dashboard_components:
        p = doc.add_paragraph()
        runner = p.add_run(f'{component}: ')
        runner.bold = True
        runner.font.color.rgb = RGBColor(212, 175, 55)
        p.add_run(desc)
    
    add_note_box(doc, "Dashboard diupdate secara real-time. Refresh halaman untuk melihat data terbaru.")
    
    doc.add_page_break()
    
    # ========== CHAPTER 4: TRANSAKSI ==========
    doc.add_heading('4. MANAJEMEN TRANSAKSI', level=1)
    
    doc.add_heading('4.1 Membuat Transaksi Baru', level=2)
    
    doc.add_paragraph("Berikut langkah detail untuk membuat transaksi:")
    doc.add_paragraph()
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Halaman Transaksi",
        "Klik menu 'Transaksi' di sidebar kiri.",
        "Halaman transaksi akan terbuka menampilkan daftar transaksi hari ini. Di bagian atas ada tombol navigasi tanggal dan tombol '+Transaksi Baru' berwarna kuning di pojok kanan."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Klik Tombol Transaksi Baru",
        "Klik tombol kuning '+ Transaksi Baru' di pojok kanan atas.",
        "Dialog form transaksi akan muncul di tengah layar. Form ini memiliki background gelap dengan border emas. Judul 'Transaksi Baru' terlihat di bagian atas dialog."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Pilih atau Tambah Nasabah",
        "Di bagian atas form, ada dropdown 'Pilih Nasabah'. Klik untuk membuka daftar nasabah, atau ketik untuk mencari.",
        "Dropdown akan menampilkan kode nasabah (contoh: MBA92B520D) dan nama. Jika nasabah baru, klik tombol '+ Nasabah Baru' di bawah dropdown untuk membuka form pendaftaran nasabah."
    )
    
    add_note_box(doc, "Sistem akan otomatis generate kode nasabah format MBA + 8 angka random untuk nasabah baru.")
    
    add_step_with_screenshot(
        doc, 4,
        "Pilih Tipe Transaksi",
        "Pilih tipe transaksi: 'Jual' (Anda menjual valas ke nasabah) atau 'Beli' (Anda membeli valas dari nasabah).",
        "Ada 2 tombol toggle: 'Jual' berwarna hijau dengan icon panah atas, 'Beli' berwarna biru dengan icon panah bawah. Tombol yang dipilih akan highlight."
    )
    
    add_step_with_screenshot(
        doc, 5,
        "Pilih Mata Uang",
        "Klik dropdown 'Mata Uang' dan pilih mata uang yang akan ditransaksikan (USD, EUR, GBP, JPY, dll).",
        "Dropdown menampilkan kode mata uang (USD, EUR) dengan nama lengkap (US Dollar, Euro). Cari dengan mengetik kode atau nama."
    )
    
    add_step_with_screenshot(
        doc, 6,
        "Masukkan Jumlah",
        "Ketik jumlah valas yang akan ditransaksikan di field 'Jumlah'.",
        "Field jumlah hanya menerima angka. Sistem akan auto-format dengan separator ribuan saat Anda mengetik."
    )
    
    add_step_with_screenshot(
        doc, 7,
        "Masukkan Kurs (Exchange Rate)",
        "Ketik kurs yang berlaku di field 'Kurs'. Kurs adalah nilai tukar 1 unit valas dalam Rupiah.",
        "Field kurs auto-format dengan separator ribuan. Contoh: untuk USD 15.500 berarti 1 USD = Rp 15.500"
    )
    
    add_step_with_screenshot(
        doc, 8,
        "Cek Total IDR",
        "Sistem akan otomatis menghitung Total IDR = Jumlah × Kurs.",
        "Total IDR ditampilkan di bawah field kurs dengan format Rupiah lengkap. Contoh: Rp 1.550.000. Total ini otomatis update saat Anda ubah jumlah atau kurs."
    )
    
    add_step_with_screenshot(
        doc, 9,
        "Isi Informasi Tambahan (Opsional)",
        "Isi field opsional: No. Voucher, Delivery Channel, Payment Method, Notes, Transaction Purpose.",
        "Field ini tidak wajib diisi. No. Voucher akan tampil sebagai '-' di daftar transaksi jika dikosongkan."
    )
    
    add_step_with_screenshot(
        doc, 10,
        "Minta Tanda Tangan Nasabah",
        "Gulir ke bawah dan minta nasabah untuk tanda tangan di area Signature Pad (kotak putih).",
        "Area signature pad adalah kotak putih kosong. Nasabah bisa tanda tangan dengan mouse, touchpad, atau touchscreen. Ada tombol 'Clear' untuk menghapus jika salah."
    )
    
    add_warning_box(doc, "Transaksi HARUS ada tanda tangan nasabah sesuai regulasi Bank Indonesia!")
    
    add_step_with_screenshot(
        doc, 11,
        "Simpan Transaksi",
        "Klik tombol 'Cetak & Simpan' untuk menyimpan dan langsung mencetak nota transaksi.",
        "Dialog akan tertutup dan nota transaksi akan terbuka di tab baru untuk dicetak. Transaksi otomatis tersimpan dan muncul di daftar transaksi."
    )
    
    doc.add_page_break()
    
    # ========== MULTI-CURRENCY TRANSACTION ==========
    doc.add_heading('4.2 Transaksi Multi-Currency', level=2)
    
    doc.add_paragraph(
        "Untuk nasabah yang menukar lebih dari 1 mata uang sekaligus (misalnya USD + EUR + GBP), "
        "gunakan fitur Multi-Currency Transaction:"
    )
    doc.add_paragraph()
    
    add_step_with_screenshot(
        doc, 1,
        "Aktifkan Mode Multi-Currency",
        "Di form transaksi, centang checkbox 'Multi-Currency Transaction' di bagian atas.",
        "Setelah dicentang, form akan berubah. Field mata uang, jumlah, dan kurs sekarang dalam bentuk tabel dengan beberapa baris."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Tambah Baris Mata Uang",
        "Klik tombol '+ Tambah Mata Uang' untuk menambah baris mata uang baru.",
        "Setiap baris memiliki: dropdown Tipe (Jual/Beli), dropdown Mata Uang, field Jumlah, field Kurs. Anda bisa tambah unlimited baris."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Isi Detail Setiap Mata Uang",
        "Untuk setiap baris, pilih tipe, mata uang, masukkan jumlah dan kurs.",
        "Setiap baris bisa punya tipe berbeda (misal baris 1 JUAL USD, baris 2 BELI EUR). Total per baris dihitung otomatis."
    )
    
    add_step_with_screenshot(
        doc, 4,
        "Lihat Grand Total",
        "Sistem menghitung Grand Total dari semua baris mata uang.",
        "Grand Total ditampilkan dengan format besar dan bold di bawah tabel. Ini adalah total keseluruhan transaksi dalam Rupiah."
    )
    
    add_step_with_screenshot(
        doc, 5,
        "Simpan Multi-Currency Transaction",
        "Isi No. Voucher yang SAMA untuk semua mata uang (penting!), minta tanda tangan, dan klik 'Simpan'.",
        "Sistem akan membuat beberapa transaksi dengan nomor transaksi berbeda tapi voucher sama. Nomor transaksi akan ada suffix -a, -b, -c untuk membedakan."
    )
    
    add_warning_box(doc, "PENTING: Untuk multi-currency, pastikan No. Voucher DIISI dan SAMA untuk semua mata uang agar bisa ditrack sebagai 1 transaksi!")
    
    doc.add_page_break()
    
    # ========== CHAPTER 5: DATA NASABAH ==========
    doc.add_heading('5. DATA NASABAH', level=1)
    
    doc.add_heading('5.1 Menambah Nasabah Baru', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Halaman Data Nasabah",
        "Klik menu 'Data Nasabah' di sidebar kiri.",
        "Halaman Data Nasabah menampilkan tabel nasabah dengan kolom: Kode, Jenis, Nama, JK, No. Identitas, Telepon, Pekerjaan, Alamat, dan Aksi."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Klik Tambah Nasabah",
        "Klik tombol '+ Tambah Nasabah' berwarna kuning di pojok kanan atas.",
        "Dialog form nasabah akan muncul. Ada 2 tab di atas: 'Perorangan' dan 'Badan Usaha'. Pilih sesuai jenis nasabah."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Isi Data Nasabah Perorangan",
        "Untuk nasabah perorangan, isi: Nama, Gender, Jenis Identitas (KTP/SIM/Passport), No. Identitas, Tempat & Tanggal Lahir, Alamat, Telepon, Pekerjaan, Sumber Dana.",
        "Form memiliki 2 kolom. Field bertanda * (bintang merah) adalah wajib diisi. Field lain opsional tapi sebaiknya dilengkapi untuk KYC."
    )
    
    add_step_with_screenshot(
        doc, 4,
        "Isi Data Nasabah Badan Usaha",
        "Untuk badan usaha, klik tab 'Badan Usaha' dan isi: Jenis Badan Usaha, Nama, No. Izin, NPWP, Alamat, PIC (Person in Charge) detail.",
        "Tab Badan Usaha memiliki field berbeda: entity_type, license_number, NPWP, dll. Semua sesuai requirement regulasi BI untuk badan usaha."
    )
    
    add_step_with_screenshot(
        doc, 5,
        "Simpan Data Nasabah",
        "Setelah semua data terisi, klik tombol 'Simpan'.",
        "Notifikasi 'Nasabah berhasil ditambahkan' akan muncul. Dialog tertutup dan nasabah baru muncul di tabel dengan kode MBA yang auto-generated."
    )
    
    doc.add_heading('5.2 Melihat Profil Nasabah', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Klik Icon Mata (View)",
        "Dari daftar nasabah, klik icon mata di kolom Aksi untuk melihat profil lengkap nasabah.",
        "Dialog profil nasabah akan muncul dengan 3 tab: V-Card Member, Buku Transaksi, dan KYC."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Lihat Member Card",
        "Tab V-Card Member menampilkan kartu member nasabah dengan QR code.",
        "Kartu member menampilkan: kode nasabah, nama, nomor identitas, alamat, dan QR code untuk scan. Ada tombol 'Print Member Card' untuk mencetak."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Lihat Buku Transaksi",
        "Klik tab 'Buku Transaksi', lalu klik tombol 'Muat Transaksi' untuk melihat riwayat transaksi nasabah.",
        "Tombol 'Muat Transaksi' akan loading 2-3 detik, lalu muncul tabel transaksi lengkap dengan total pembelian, penjualan, dan grand total. Ada tombol 'Cetak Buku Transaksi' untuk print."
    )
    
    add_note_box(doc, "Fitur lazy loading untuk Buku Transaksi membuat halaman profil nasabah load lebih cepat (3-5x). Data transaksi hanya dimuat saat Anda klik tombol 'Muat Transaksi'.")
    
    add_step_with_screenshot(
        doc, 4,
        "Lihat Data KYC",
        "Tab 'KYC' menampilkan data lengkap nasabah untuk keperluan audit dan compliance.",
        "KYC tab berisi semua data nasabah yang terstruktur: identitas, alamat, pekerjaan, sumber dana, dll. Ada tombol 'Print KYC' untuk cetak form KYC sesuai format BI."
    )
    
    doc.add_page_break()
    
    # ========== CHAPTER 6: BUKU KAS ==========
    doc.add_heading('6. BUKU KAS', level=1)
    
    doc.add_paragraph(
        "Buku Kas mencatat semua aliran kas harian (masuk/keluar). "
        "System otomatis mencatat entry buku kas dari setiap transaksi."
    )
    
    doc.add_heading('6.1 Melihat Buku Kas', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Halaman Buku Kas",
        "Klik menu 'Buku Kas' di sidebar.",
        "Halaman Buku Kas menampilkan 4 kartu summary di atas: Saldo Awal, Total Debit, Total Kredit, Saldo Akhir. Di bawahnya ada tabel entry kas dengan kolom: Tanggal, Tipe, Keterangan, Debit, Kredit, Sumber, Aksi."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Navigasi Tanggal",
        "Gunakan tombol panah kiri/kanan atau klik icon kalender untuk memilih tanggal.",
        "Ada 3 tombol navigasi: '← Sebelumnya', date picker (dengan icon kalender), '→ Berikutnya', dan tombol 'Hari Ini' untuk kembali ke tanggal sekarang."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Filter Data (Opsional)",
        "Klik tombol 'Filter' untuk filter berdasarkan tipe entry, sumber, atau range tanggal.",
        "Panel filter akan muncul dengan dropdown: Tipe Entry (Semua/Debit/Kredit), Tipe Transaksi (Pembelian/Penjualan), Sumber (Semua/Manual/Transaksi), dan range tanggal."
    )
    
    doc.add_heading('6.2 Menambah Entry Manual', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Klik Tambah Entri",
        "Klik tombol '+ Tambah Entri' berwarna kuning.",
        "Dialog form entry manual akan muncul dengan field: Cabang, Tipe (Debit/Kredit), Jumlah, dan Keterangan."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Isi Detail Entry",
        "Pilih cabang, pilih tipe (Debit untuk kas masuk, Kredit untuk kas keluar), masukkan jumlah dan keterangan.",
        "Keterangan bisa diisi bebas, contoh: 'Biaya operasional', 'Kas awal hari', 'Setoran bank', dll."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Simpan Entry Manual",
        "Klik tombol 'Simpan'.",
        "Entry manual akan muncul di tabel dengan label 'Manual' di kolom Sumber. Entry manual hanya bisa di-edit/delete oleh Admin."
    )
    
    doc.add_heading('6.3 Tools Diagnostik Buku Kas (Admin Only)', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Klik Tombol Periksa Data",
        "Di halaman Buku Kas, klik tombol kuning 'Periksa Data' (ada icon peringatan).",
        "Dialog 'Hasil Pemeriksaan Data' akan muncul menampilkan statistik: Total Transaksi, Total Buku Kas, dan Ketidaksesuaian (jika ada)."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Review Hasil Pemeriksaan",
        "Jika ada ketidaksesuaian, akan ditampilkan detail: nomor transaksi, nilai di Transaksi vs Buku Kas, dan selisihnya.",
        "Status 'Data Konsisten!' berarti tidak ada masalah. Status '⚠ Ditemukan Ketidaksesuaian!' berarti ada transaksi yang tidak sinkron."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Perbaiki Data (Jika Ada Masalah)",
        "Jika ditemukan ketidaksesuaian, ada 2 opsi: '🔧 Perbaiki Sekarang' (soft fix) atau '♻️ Hitung Ulang (Advanced)' (hard fix/recreate).",
        "'Perbaiki Sekarang' akan update nilai yang salah. 'Hitung Ulang' akan delete dan recreate semua entry dari data transaksi asli (lebih aman untuk data yang sangat rusak)."
    )
    
    add_warning_box(doc, "Hitung Ulang (Advanced) akan menghapus dan membuat ulang entry buku kas. Gunakan hanya jika data sudah sangat rusak dan 'Perbaiki Sekarang' tidak cukup.")
    
    doc.add_page_break()
    
    # ========== CHAPTER 7: MUTASI VALAS ==========
    doc.add_heading('7. MUTASI VALAS', level=1)
    
    doc.add_paragraph(
        "Mutasi Valas menampilkan pergerakan stok setiap mata uang per hari. "
        "Halaman ini penting untuk monitoring stock dan menghitung profit/loss."
    )
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Halaman Mutasi Valas",
        "Klik menu 'Mutasi Valas' di sidebar.",
        "Halaman menampilkan tabel mutasi per mata uang dengan kolom: Tanggal, Mata Uang, Stock Awal, Pembelian, Penjualan, Stock Akhir, Avg Rate, Profit/Loss."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Pilih Tanggal dan Cabang",
        "Gunakan navigasi tanggal dan dropdown cabang untuk filter data yang ingin dilihat.",
        "Data mutasi otomatis dihitung dari transaksi. Stock Akhir hari ini = Stock Awal hari besok (dijaga otomatis oleh sistem snapshot)."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Export Data Mutasi",
        "Klik tombol 'Excel' atau 'PDF' untuk export data mutasi.",
        "File export berisi detail lengkap mutasi per mata uang untuk periode yang dipilih."
    )
    
    add_note_box(doc, "Jika ada ketidaksesuaian stock (Stock Akhir hari X ≠ Stock Awal hari X+1), gunakan tombol 'Perbaiki Data' yang tersedia untuk Admin.")
    
    doc.add_page_break()
    
    # ========== CHAPTER 8: LAPORAN ==========
    doc.add_heading('8. LAPORAN', level=1)
    
    doc.add_heading('8.1 Laporan Transaksi', level=2)
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Halaman Laporan",
        "Klik menu 'Laporan' di sidebar.",
        "Ada 2 tab: 'Laporan Transaksi' dan 'SIPESAT'. Default terbuka di tab Laporan Transaksi."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Pilih Periode Laporan",
        "Isi Tanggal Mulai dan Tanggal Akhir, lalu pilih Cabang (atau pilih 'Semua Cabang').",
        "Date picker memudahkan pemilihan tanggal. Pilih periode sesuai kebutuhan (harian, mingguan, bulanan, atau custom range)."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Generate dan Download Laporan",
        "Klik tombol 'Buat Laporan'. System akan generate laporan dan langsung download.",
        "Laporan berisi ringkasan dan detail transaksi periode yang dipilih. Format Excel untuk analisa lebih lanjut."
    )
    
    doc.add_heading('8.2 Laporan SIPESAT', level=2)
    
    doc.add_paragraph(
        "SIPESAT adalah sistem pelaporan ke Bank Indonesia. Laporan ini wajib dibuat dan disubmit secara berkala."
    )
    
    add_step_with_screenshot(
        doc, 1,
        "Buka Tab SIPESAT",
        "Klik tab 'SIPESAT' di halaman Laporan.",
        "Form SIPESAT memiliki field lengkap sesuai requirement BI: ID Pengguna Jasa Keuangan, periode, dll."
    )
    
    add_step_with_screenshot(
        doc, 2,
        "Isi Detail SIPESAT",
        "Isi ID PJK, periode laporan, dan informasi yang diminta.",
        "Pastikan ID PJK sudah terdaftar di Pengaturan > Perusahaan > IDPJK."
    )
    
    add_step_with_screenshot(
        doc, 3,
        "Generate File SIPESAT",
        "Klik 'Generate SIPESAT' untuk membuat file sesuai format BI.",
        "File akan ter-download dalam format yang diminta BI (.txt atau .csv). Upload file ini ke portal SIPESAT BI."
    )
    
    doc.add_page_break()
    
    # ========== CHAPTER 9: PENGATURAN ==========
    doc.add_heading('9. PENGATURAN (ADMIN ONLY)', level=1)
    
    doc.add_paragraph("Halaman Pengaturan hanya dapat diakses oleh Admin. Berisi 7 tab:")
    
    doc.add_heading('9.1 Tab Perusahaan', level=2)
    doc.add_paragraph("Update profil perusahaan: nama, alamat, telepon, email, website, nomor izin BI, NPWP, IDPJK.")
    
    doc.add_heading('9.2 Tab Pengguna', level=2)
    doc.add_paragraph("Manajemen user: tambah, edit, hapus user. Set role (Admin/Kasir/Teller) dan assign cabang.")
    
    doc.add_heading('9.3 Tab Cabang', level=2)
    doc.add_paragraph("Manajemen cabang: tambah, edit cabang. Set saldo awal untuk setiap mata uang per cabang.")
    
    doc.add_heading('9.4 Tab Mata Uang', level=2)
    doc.add_paragraph("Manajemen mata uang: tambah, edit mata uang yang tersedia untuk transaksi.")
    
    doc.add_heading('9.5 Tab Log Aktivitas', level=2)
    doc.add_paragraph("Audit trail: melihat semua aktivitas user (login, logout, create transaction, dll).")
    
    doc.add_heading('9.6 Tab Petunjuk Teknis', level=2)
    doc.add_paragraph("Download petunjuk teknis ini dalam format Word (.docx) atau PDF (.pdf).")
    
    doc.add_heading('9.7 Tab Maintenance', level=2)
    doc.add_paragraph(
        "Tools maintenance untuk Admin:\n"
        "• Migrasi Format Tanggal - untuk normalisasi format date di database\n"
        "• Periksa Status Data - dry-run simulation untuk melihat berapa records perlu update\n"
        "• Jalankan Migrasi - execute migrasi sebenarnya (backup database dulu!)"
    )
    
    doc.add_page_break()
    
    # ========== CHAPTER 10: EXPORT DATA ==========
    doc.add_heading('10. BACKUP DAN EXPORT DATA', level=1)
    
    doc.add_heading('10.1 Export ke Excel', level=2)
    doc.add_paragraph(
        "Hampir semua halaman memiliki tombol 'Excel'. Klik untuk export data halaman tersebut ke format Excel (.xlsx). "
        "Data akan ter-download otomatis dan bisa dibuka di Microsoft Excel atau Google Sheets."
    )
    
    doc.add_heading('10.2 Export ke PDF', level=2)
    doc.add_paragraph(
        "Tombol 'PDF' akan export data dalam format PDF yang siap untuk print atau archive. "
        "PDF sudah terformat rapi dengan header perusahaan."
    )
    
    doc.add_heading('10.3 Cetak Langsung', level=2)
    doc.add_paragraph(
        "Tombol 'Cetak' akan membuka print preview browser. "
        "Anda bisa langsung print atau save as PDF dari browser."
    )
    
    doc.add_page_break()
    
    # ========== CHAPTER 11: TIPS & TROUBLESHOOTING ==========
    doc.add_heading('11. TIPS DAN TROUBLESHOOTING', level=1)
    
    doc.add_heading('11.1 Tips Penggunaan', level=2)
    
    tips = [
        ('Gunakan Search di Transaksi', 'Ketik nomor transaksi, kode nasabah, atau nama untuk cari transaksi cepat.'),
        ('Nomor Voucher Kosong = "-"', 'Jika voucher tidak diisi, akan tampil "-" di daftar transaksi (bukan kosong).'),
        ('Multi-Currency = Voucher Wajib', 'Untuk transaksi multi-currency, pastikan isi No. Voucher yang sama untuk semua mata uang.'),
        ('Lazy Load Transaksi Nasabah', 'Klik "Muat Transaksi" di profil nasabah untuk performa lebih baik.'),
        ('Backup Rutin', 'Export data ke Excel secara berkala sebagai backup.'),
    ]
    
    for title, desc in tips:
        p = doc.add_paragraph()
        runner = p.add_run(f'✅ {title}: ')
        runner.bold = True
        p.add_run(desc)
    
    doc.add_heading('11.2 Troubleshooting Umum', level=2)
    
    issues = [
        ('Data Buku Kas Tidak Match dengan Transaksi', 
         'Gunakan tools "Periksa Data" di halaman Buku Kas. Jika ada ketidaksesuaian, gunakan "Hitung Ulang (Advanced)" untuk fix.'),
        ('Stock Mutasi Valas Tidak Konsisten',
         'Gunakan tombol "Perbaiki Data" di halaman Mutasi Valas (Admin only). System akan recalculate semua snapshot.'),
        ('Transaksi Tidak Muncul di Daftar',
         'Pastikan tanggal yang dipilih benar. Transaksi disimpan dengan timezone WITA (UTC+8). Cek juga filter yang aktif.'),
        ('Tidak Bisa Login',
         'Pastikan email dan password benar (case-sensitive). Jika lupa password, hubungi Administrator.'),
        ('Halaman Lambat',
         'Kurangi range tanggal yang dipilih. Gunakan filter untuk limit data yang dimuat. Fitur lazy loading sudah diterapkan di profil nasabah.'),
    ]
    
    for issue, solution in issues:
        p = doc.add_paragraph()
        runner = p.add_run(f'❓ {issue}\n')
        runner.bold = True
        runner.font.color.rgb = RGBColor(251, 191, 36)
        runner = p.add_run(f'💡 Solusi: {solution}')
        runner.font.color.rgb = RGBColor(110, 231, 183)
        doc.add_paragraph()
    
    doc.add_heading('11.3 Kontak Support', level=2)
    doc.add_paragraph(
        "Jika mengalami kendala teknis yang tidak bisa diselesaikan dengan panduan ini, "
        "hubungi Administrator atau IT Support perusahaan Anda."
    )
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = footer_para.add_run('\n\n--- AKHIR DOKUMEN ---\n\n')
    runner.bold = True
    runner = footer_para.add_run(f'Mulia Bali Valuta (MBA) Money Changer\n{datetime.now().year}')
    runner.italic = True
    
    # Save document
    doc.save(output_path)
    return output_path


def create_enhanced_user_manual_pdf(output_path: str):
    """Create enhanced PDF manual with detailed instructions"""
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#D4AF37'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#D4AF37'),
        spaceBefore=12,
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.HexColor('#6EE7B7'),
        spaceBefore=10,
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        leading=14,
        spaceAfter=6,
        alignment=TA_JUSTIFY
    )
    
    note_style = ParagraphStyle(
        'NoteStyle',
        parent=styles['BodyText'],
        fontSize=9,
        textColor=colors.HexColor('#6EE7B7'),
        leftIndent=20,
        rightIndent=20,
        spaceAfter=10,
        borderPadding=5
    )
    
    # Title Page
    story.append(Paragraph("PETUNJUK TEKNIS PENGGUNAAN", title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("APLIKASI MULIA BALI VALUTA", heading1_style))
    story.append(Paragraph("Money Changer Management System", body_style))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(f"Versi 2.0 - Edisi Lengkap<br/>{datetime.now().strftime('%B %Y')}", body_style))
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph("DAFTAR ISI", heading1_style))
    story.append(Spacer(1, 0.2*inch))
    
    toc_data = [
        ['1. Pendahuluan', ''],
        ['2. Login dan Logout', ''],
        ['3. Dashboard', ''],
        ['4. Manajemen Transaksi', ''],
        ['5. Data Nasabah', ''],
        ['6. Buku Kas', ''],
        ['7. Mutasi Valas', ''],
        ['8. Laporan', ''],
        ['9. Pengaturan', ''],
        ['10. Backup dan Export', ''],
        ['11. Tips dan Troubleshooting', ''],
    ]
    
    toc_table = Table(toc_data, colWidths=[5*inch, 0.5*inch])
    toc_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#FEF3C7')),
        ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor('#D4AF37')),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(toc_table)
    story.append(PageBreak())
    
    # Content chapters (simplified for PDF)
    story.append(Paragraph("1. PENDAHULUAN", heading1_style))
    story.append(Paragraph(
        "Mulia Bali Valuta (MBA) Money Changer adalah aplikasi manajemen transaksi "
        "penukaran valuta asing dengan fitur lengkap untuk operasional modern.",
        body_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Add similar content as DOCX...
    # (Simplified for now - can be expanded based on user feedback)
    
    story.append(Paragraph(
        "📌 Catatan: Untuk panduan lengkap dengan step-by-step detail, lihat versi Word (.docx) dari petunjuk teknis ini.",
        note_style
    ))
    
    # Build PDF
    doc.build(story)
    return output_path
